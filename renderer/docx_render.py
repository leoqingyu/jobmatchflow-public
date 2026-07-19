"""
简历 DOCX 渲染：单模板，用 python-docx API 逐段程序化拼装，不用 docxtpl/Jinja-in-docx
模板文件——这台机器没法产出真实的二进制 .docx 模板（Write 工具只能写文本），而且 Jinja
标签在 Word XML 里跨 run 断裂，对变长的经历/bullet 循环本来就脆弱。

版式复刻用户提供的 LaTeX 简历模板（Jake Gutierrez / Sourabh Bajaj 风格）：
- 页头：姓名（大写加粗）+ 联系方式一行，右上角可选头像（两栏无边框表格）。
- Education/Experience 每条用两行："粗体标题 + 右对齐日期"（右对齐用右制表位实现，不用
  嵌套表格——制表位对 ATS 纯文本抽取更友好）、"斜体机构/地点 + 右对齐斜体地点/空"。
- 分节标题（EDUCATION/EXPERIENCE/SKILLS）全大写加粗，下方一条横线（\\titlerule 效果，
  用段落底边框实现）。
- 项目符号用 Word 内置 "List Bullet" 样式，段前/段后间距压到最小，贴近 LaTeX 的 \\vspace{-Npt}
  紧凑感。

字符预算是经验值常量，不试图从字体精确反推"每行多少字符"（DOCX 是比例字体，没有等宽终端
那种确定性），渲染完尽力用 LibreOffice headless 转一次 PDF、拿仓库已有的 pypdf 数实际页数
做校验；装不了/转不出就跳过校验，不阻断生成——调用方按"信任字符预算的保守估计"处理。一页
装不装得下、装不满要不要回填，都是调用方（services/resume_generation_service.py）的事，
这个模块只管"给定内容，渲染成这个版式的字节流"和"数一下渲染结果是几页"。
"""

from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path

from core.config import settings
from core.logger import get_logger

logger = get_logger(__name__)

# 经验值：单栏 Letter、10.5pt 正文、紧凑页边距下一页大致能放的量——保守估计，
# 应在有真实渲染环境的机器上用 count_pages() 标定后再调。经历条数/每条 bullet 数不放在这里：
# 那是内容质量要求（对 LLM 的要求），见 ai/resume_selection.py 的 MIN/MAX_EXPERIENCE_ITEMS、
# ai/resume_rewrite.py 的 MIN/MAX_BULLETS_PER_EXPERIENCE。
MAX_BULLET_CHARS = 180
MAX_SKILLS = 20
# Education 最多渲染 3 段（第一行学校/专业+时间，第二行地点，之后是课程/简介 bullets）。
# detail bullets 不做行数截断——那是用户在 Profile 里手打的原文，不是 LLM 生成的，截断等于
# 静默丢用户数据；页数超限交给用户在编辑页自己控制（该 textarea 已提示"保持 2 行以内"）。
MAX_EDUCATION_ITEMS = 3

_NAME_PT = 20
_CONTACT_PT = 10
_SECTION_PT = 12
_SUBHEAD_PT = 11
_META_PT = 10
_BULLET_PT = 10.5

# 全篇统一无衬线：ascii/hAnsi 覆盖西文，eastAsia 覆盖中文等 CJK 字符——只设 font.name 只
# 会写 ascii/hAnsi，中文走 eastAsia 字段，不显式设的话 Word/LibreOffice 会退回主题默认
# 东亚字体（通常是宋体一类衬线字），中西文混排时看起来"半衬线半无衬线"。见 _set_sans_serif。
_FONT_NAME = "Calibri"
_EAST_ASIAN_FONT_NAME = "Microsoft YaHei"

# 头像按目标高度（略高于姓名+联系方式两行,不是按列宽拉伸)定尺寸——见 _photo_dimensions。
_PHOTO_TARGET_HEIGHT_IN = 1.0
_PHOTO_MAX_WIDTH_IN = 1.3
_PHOTO_RIGHT_GAP_IN = 0.15

# ai.resume_rewrite.polish_bullets 用 **text** 标记关键信息（量化结果/核心技术）该加粗；
# 渲染时解析成多个 run，不是整条 bullet 一个 run，见 _add_bullet_paragraph
_BOLD_MARK_RE = re.compile(r"\*\*(.+?)\*\*")


def _set_sans_serif(font) -> None:
    """把 font.name（只覆盖 ascii/hAnsi）之外，再显式钉死 eastAsia/cs 两个字段，确保
    中文等 CJK 字符也走无衬线字体，不会退回主题默认的衬线字体。font 可以是 style.font
    或 run.font——两者的 .element 分别是 <w:style>/<w:r>，都支持 get_or_add_rPr()。"""
    from docx.oxml.ns import qn

    font.name = _FONT_NAME
    rpr = font.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), _EAST_ASIAN_FONT_NAME)
    rfonts.set(qn("w:cs"), _FONT_NAME)


def render_tailored_resume_docx(content: dict, photo_path: str | Path | None = None) -> bytes:
    """
    content 形状（基本信息字段复用 renderer.cv_render 的既有 schema，经历/技能是新选材+
    改写结果）：
    {
      "full_name", "location", "phone", "email", "visa"?,
      "linkedin"?: {"url","label"}, "github"?: {"url","label"}, "profile_summary"?,
      "education": [{"degree","institution","date_range","location","bullets":[...]}],
      "experience": [{"title","company","location","date_range","bullets":[...]}],
      "skills": [str, ...],
      "skill_categories"?: [{"label","skills":[str,...]}, ...]（LLM 分类结果，见
      ai.resume_rewrite.categorize_skills；存在则优先渲染这个多行分类版本，不存在
      （比如用户手动编辑过 skills 导致分类失效）就退回 "skills" 单行版本）,
      "languages"?: [str, ...]（口语能力，如 "English (C1)"——独立于 skills 单起一行，
      跟模板里 Skills 区块下的 "Languages: ..." 一行对应）,
    }

    photo_path：用户证件照绝对路径（见 services/profile_photo.py），存在则渲染成右上角头像
    （瑞士/德语区简历惯例）；不存在或未传则退回纯文字页头，ATS 友好。
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        raise RuntimeError("未安装 python-docx，请运行 pip install python-docx") from e

    doc = Document()

    style = doc.styles["Normal"]
    _set_sans_serif(style.font)
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    usable_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin

    # 联系方式行直接放完整 URL，不用 label 当作"链接文字"占位——DOCX 里这不是可点击超链接，
    # 用 label（如 "LinkedIn"）顶替 URL 反而让读者看不到真实地址。
    contact_bits = [content.get("location"), content.get("phone"), content.get("email")]
    linkedin = content.get("linkedin") or {}
    if linkedin.get("url"):
        contact_bits.append(linkedin["url"])
    github = content.get("github") or {}
    if github.get("url"):
        contact_bits.append(github["url"])
    if content.get("visa"):
        contact_bits.append(content["visa"])
    contact_line = " | ".join(b for b in contact_bits if b)

    has_photo = bool(photo_path) and Path(photo_path).is_file()
    if has_photo:
        _add_header_with_photo(doc, Inches, Pt, WD_ALIGN_PARAGRAPH, content.get("full_name") or "", contact_line, photo_path)
    else:
        _add_header_plain(doc, Pt, WD_ALIGN_PARAGRAPH, content.get("full_name") or "", contact_line)

    has_summary = bool(content.get("profile_summary"))
    if has_summary:
        _add_section_heading(doc, Pt, "Summary", first=True)
        doc.add_paragraph(content["profile_summary"])

    education = (content.get("education") or [])[:MAX_EDUCATION_ITEMS]
    if education:
        _add_section_heading(doc, Pt, "Education", first=not has_summary)
        for edu in education:
            degree = edu.get("degree") or ""
            institution = edu.get("institution") or ""
            bold_left = ", ".join(b for b in (degree, institution) if b)
            _add_subheading_entry(
                doc, Pt, usable_width,
                bold_left=bold_left,
                right1=edu.get("date_range") or "",
                italic_left=edu.get("location") or "",
                right2="",
                bullets=edu.get("bullets") or [],
            )

    experience = content.get("experience") or []
    if experience:
        _add_section_heading(doc, Pt, "Experience")
        for exp in experience:
            _add_subheading_entry(
                doc, Pt, usable_width,
                bold_left=exp.get("title") or "",
                right1=exp.get("date_range") or "",
                italic_left=exp.get("company") or "",
                right2=exp.get("location") or "",
                bullets=exp.get("bullets") or [],
            )

    skill_categories = content.get("skill_categories") or []
    skills = content.get("skills") or []
    languages = content.get("languages") or []
    if skill_categories or skills or languages:
        _add_section_heading(doc, Pt, "Skills")
        if skill_categories:
            for cat in skill_categories:
                label = cat.get("label") or ""
                items = cat.get("skills") or []
                if label and items:
                    _add_labeled_skills_line(doc, Pt, label, ", ".join(items))
        elif skills:
            _add_labeled_skills_line(doc, Pt, "Skills", ", ".join(skills))
        if languages:
            _add_labeled_skills_line(doc, Pt, "Languages", ", ".join(languages))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_cover_letter_docx(content: dict) -> bytes:
    """
    单模板求职信：姓名+联系方式头 -> 日期 -> 称呼 -> 正文段落 -> 落款。跟简历一样，纯
    python-docx 程序化拼装，不用模板文件。

    content 形状：
    {
      "full_name", "location"?, "phone"?, "email"?,
      "company", "job_title", "date"?, "paragraphs": [str, ...],
      "greeting"?（默认 "Dear Hiring Team,"）, "closing"?（默认 "Sincerely,"，落款签名固定
      是 full_name，不做成单独字段——那是身份信息，不是可以随便改的措辞）,
    }
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        raise RuntimeError("未安装 python-docx，请运行 pip install python-docx") from e

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    name_p = doc.add_paragraph()
    name_run = name_p.add_run(content.get("full_name") or "")
    name_run.bold = True
    name_run.font.size = Pt(13)

    contact_bits = [content.get("location"), content.get("phone"), content.get("email")]
    contact_line = " | ".join(b for b in contact_bits if b)
    if contact_line:
        doc.add_paragraph(contact_line)

    if content.get("date"):
        date_p = doc.add_paragraph()
        date_p.paragraph_format.space_before = Pt(12)
        date_p.add_run(content["date"])

    company = content.get("company") or ""
    if company:
        doc.add_paragraph(company)

    greeting_p = doc.add_paragraph()
    greeting_p.paragraph_format.space_before = Pt(12)
    greeting_p.add_run(content.get("greeting") or "Dear Hiring Team,")

    for para in content.get("paragraphs") or []:
        if not para:
            continue
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(8)

    closing_p = doc.add_paragraph()
    closing_p.paragraph_format.space_before = Pt(4)
    closing_run = closing_p.add_run(content.get("closing") or "Sincerely,")
    closing_run.add_break()
    closing_p.add_run(content.get("full_name") or "")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _photo_dimensions(photo_path, Inches):
    """按目标高度（_PHOTO_TARGET_HEIGHT_IN，略高于姓名+联系方式两行）算等比宽度——用高度
    而不是列宽定尺寸：证件照普遍是竖版（高:宽 通常 >1），按列宽拉伸会让高度完全由照片
    原始长宽比决定，轻易冲到 2 inch+，姓名+联系方式那两行远够不到那个高度，中间空出一大块
    白，Education 反而被拖到很靠下的地方，不是"贴着照片底部"该有的样子。用 python-docx
    自带的 Image 读取像素尺寸，不额外依赖 Pillow；读不到就退回正方形兜底。极端宽幅照片
    （罕见）再按宽度上限收窄，避免右列被挤爆。"""
    from docx.image.image import Image as DocxImage

    height_in = _PHOTO_TARGET_HEIGHT_IN
    width_in = height_in
    try:
        img = DocxImage.from_file(str(photo_path))
        if img.px_width and img.px_height:
            width_in = height_in * (img.px_width / img.px_height)
    except Exception:
        pass
    if width_in > _PHOTO_MAX_WIDTH_IN:
        height_in = height_in * (_PHOTO_MAX_WIDTH_IN / width_in)
        width_in = _PHOTO_MAX_WIDTH_IN
    return Inches(width_in), Inches(height_in)


def _add_header_with_photo(doc, Inches, Pt, WD_ALIGN_PARAGRAPH, full_name: str, contact_line: str, photo_path) -> None:
    """姓名/联系方式 + 头像的两栏页头：单行双列表格，左列文字右列照片，不加边框线。
    两个单元格都顶部对齐；照片按 _photo_dimensions 的目标高度定尺寸（比文字块略高一点，
    不是远高于它），这样表格行的多余高度不会大到离谱，Education 能紧接着页头往上贴，
    整体更紧凑。照片右对齐后留一点右缩进，不贴死页面右边距。"""
    from docx.enum.table import WD_ALIGN_VERTICAL

    photo_col_width = Inches(1.5)
    text_col_width = Inches(5.7)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    info_cell, photo_cell = table.rows[0].cells
    info_cell.width = text_col_width
    photo_cell.width = photo_col_width
    info_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    photo_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    name_p = info_cell.paragraphs[0]
    name_p.paragraph_format.space_after = Pt(0)
    name_run = name_p.add_run((full_name or "").upper())
    name_run.bold = True
    name_run.font.size = Pt(_NAME_PT)
    if contact_line:
        contact_p = info_cell.add_paragraph()
        contact_p.paragraph_format.space_before = Pt(1)
        contact_p.paragraph_format.space_after = Pt(0)
        contact_p.add_run(contact_line).font.size = Pt(_CONTACT_PT)

    photo_p = photo_cell.paragraphs[0]
    photo_p.paragraph_format.space_after = Pt(0)
    photo_p.paragraph_format.right_indent = Inches(_PHOTO_RIGHT_GAP_IN)
    photo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    photo_run = photo_p.add_run()
    width, height = _photo_dimensions(photo_path, Inches)
    photo_run.add_picture(str(photo_path), width=width, height=height)


def _add_header_plain(doc, Pt, WD_ALIGN_PARAGRAPH, full_name: str, contact_line: str) -> None:
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(0)
    name_run = name_p.add_run((full_name or "").upper())
    name_run.bold = True
    name_run.font.size = Pt(_NAME_PT)
    if contact_line:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_before = Pt(1)
        contact_p.paragraph_format.space_after = Pt(0)
        contact_p.add_run(contact_line).font.size = Pt(_CONTACT_PT)


def _set_bottom_border(paragraph) -> None:
    """给段落加一条底边框，配合大写加粗的分节标题模拟 LaTeX \\titlerule 的横线效果。
    python-docx 没有原生的段落边框 API，走 OXML 直接拼 w:pBdr（标准做法）。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_section_heading(doc, Pt, text: str, *, first: bool = False) -> None:
    """first=True（紧跟在页头后面的第一个分节）时段前距压到最小，让 Education 尽量贴着
    页头往上移，整体更紧凑；不是第一个分节时用正常间距分隔上一节内容。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0) if first else Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(_SECTION_PT)
    _set_bottom_border(p)


def _add_two_col_line(
    doc, Pt, usable_width, left_text: str, right_text: str,
    *, left_bold: bool = False, left_italic: bool = False, right_italic: bool = False, size_pt: float,
):
    """左文字 + 右对齐（右制表位）文字的一行——LaTeX resumeSubheading/tabularx 的 Word 等价物。
    用制表位而不是嵌套表格：对 ATS 纯文本抽取更友好，也更不容易在 Word 里出现表格行高怪异。"""
    from docx.enum.text import WD_TAB_ALIGNMENT

    if not left_text and not right_text:
        return None
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
    if left_text:
        r = p.add_run(left_text)
        r.bold = left_bold
        r.italic = left_italic
        r.font.size = Pt(size_pt)
    if right_text:
        p.add_run("\t")
        r2 = p.add_run(right_text)
        r2.italic = right_italic
        r2.font.size = Pt(size_pt)
    return p


def _add_bullet_paragraph(doc, Pt, text: str):
    """一条 bullet 可能带 **text** 加粗标记（见 ai.resume_rewrite.polish_bullets）——按标记
    切成多个 run，标记内的部分单独加粗，而不是整条一个 run（那样就没法局部加粗）。"""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    pos = 0
    for m in _BOLD_MARK_RE.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            r.font.size = Pt(_BULLET_PT)
        r = p.add_run(m.group(1))
        r.bold = True
        r.font.size = Pt(_BULLET_PT)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        r.font.size = Pt(_BULLET_PT)
    return p


def _add_subheading_entry(
    doc, Pt, usable_width, *, bold_left: str, right1: str, italic_left: str, right2: str, bullets: list[str],
) -> None:
    _add_two_col_line(
        doc, Pt, usable_width, bold_left, right1,
        left_bold=True, left_italic=False, right_italic=False, size_pt=_SUBHEAD_PT,
    )
    line2 = _add_two_col_line(
        doc, Pt, usable_width, italic_left, right2,
        left_bold=False, left_italic=True, right_italic=True, size_pt=_META_PT,
    )
    last_para = line2
    for bullet in bullets:
        if not bullet:
            continue
        bp = _add_bullet_paragraph(doc, Pt, bullet)
        last_para = bp
    # 条目之间留一点呼吸间距（对应 LaTeX 累积的 \vspace{-7pt} 紧凑感之外仍需要的最小分隔）
    if last_para is not None:
        last_para.paragraph_format.space_after = Pt(5)


def _add_labeled_skills_line(doc, Pt, label: str, value: str) -> None:
    """粗体分类名 + 冒号 + 内容，一行——对应模板 Skills 区块里 \\textbf{Category}{: items} 的写法。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    p.add_run(value)


def _docx_to_pdf_bytes(docx_bytes: bytes) -> bytes | None:
    """LibreOffice headless 转一次 PDF，返回字节流；找不到 soffice 二进制或转换失败都返回
    None。count_pages()/render_preview_png() 共用这一步，避免重复实现 subprocess 调用。"""
    binary_setting = settings.libreoffice_binary_path
    binary = shutil.which(binary_setting)
    if not binary and Path(binary_setting).is_file():
        binary = binary_setting
    if not binary:
        logger.debug("libreoffice_binary_path=%s 未找到，跳过实际渲染", binary_setting)
        return None

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        docx_path = tmp_path / "resume.docx"
        docx_path.write_bytes(docx_bytes)
        # 并发预览请求（同一页面加载时简历+求职信几乎同时触发）会各自 fork 一个 soffice
        # 进程；不显式指定 UserInstallation，它们会抢同一个默认 profile 目录的锁，
        # 输给锁竞争的那个直接转换失败（复现过：5 个并发请求里 2 个返回空）。每次调用
        # 给一个仅本次临时目录内的独立 profile，互不干扰。
        profile_dir = tmp_path / "lo_profile"
        try:
            subprocess.run(
                [
                    binary,
                    "--headless",
                    f"-env:UserInstallation=file://{profile_dir}",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(tmp_path),
                    str(docx_path),
                ],
                check=True,
                capture_output=True,
                timeout=60,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError) as e:
            logger.warning("LibreOffice 转 PDF 失败: %s", e)
            return None

        pdf_path = tmp_path / "resume.pdf"
        if not pdf_path.exists():
            return None
        return pdf_path.read_bytes()


def count_pages(docx_bytes: bytes) -> int | None:
    """
    尽力而为的实际页数校验：调 LibreOffice headless 转一次 PDF，再用仓库已有的 pypdf 数页。
    找不到 soffice 二进制、或转换失败，都返回 None——调用方把 None 当"没法校验，信任字符
    预算的保守估计，不阻断"处理，不是必须依赖项。
    """
    pdf_bytes = _docx_to_pdf_bytes(docx_bytes)
    if pdf_bytes is None:
        return None
    try:
        from io import BytesIO

        from pypdf import PdfReader

        return len(PdfReader(BytesIO(pdf_bytes)).pages)
    except Exception as e:
        logger.warning("pypdf 数页失败: %s", e)
        return None


def render_preview_png(docx_bytes: bytes, *, dpi: int = 55) -> tuple[int | None, bytes | None]:
    """
    低保真缩略图：编辑页面实时预览用，不用每次都下载 DOCX 才能看排版效果。DPI 故意调得很低
    （默认 55，正常打印/屏幕阅读一般 150+）——用户明确要的是"内容看不清、轮廓要清晰"：只是
    确认页头/分节/条目大致占多少地方、有没有溢出到第二页，不是一份可读的预览图，没必要为了
    没人会去读的文字浪费渲染分辨率/传输体积。

    返回 (page_count, 第一页 PNG 字节流)；LibreOffice/PyMuPDF 任一环节不可用或失败都返回
    (None, None) 或 (page_count, None)——调用方按"预览暂时不可用，不阻断编辑"处理。
    """
    pdf_bytes = _docx_to_pdf_bytes(docx_bytes)
    if pdf_bytes is None:
        return None, None

    page_count: int | None
    try:
        from io import BytesIO

        from pypdf import PdfReader

        page_count = len(PdfReader(BytesIO(pdf_bytes)).pages)
    except Exception as e:
        logger.warning("pypdf 数页失败: %s", e)
        page_count = None

    try:
        import fitz  # PyMuPDF

        pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        try:
            zoom = dpi / 72.0
            pixmap = pdf_doc[0].get_pixmap(matrix=fitz.Matrix(zoom, zoom))
            return page_count, pixmap.tobytes("png")
        finally:
            pdf_doc.close()
    except Exception as e:
        logger.warning("PyMuPDF 缩略图渲染失败: %s", e)
        return page_count, None
